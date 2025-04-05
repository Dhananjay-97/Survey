import os
import json
import configparser
import logging
from datetime import datetime
from typing import Any, List, Union, Tuple
import requests
import openai
import pandas as pd
import uuid
import copy
from docx import Document
import time


if not os.path.exists('logs/'):
    os.mkdir('logs/')
if not os.path.exists('questionnaires/'):
    os.mkdir('questionnaires/')

class SurveyGenerator:
    """Survey generator class for creating survey using GPT3"""
    def __init__(self) -> None:
        """Initialize the survey generator parameters."""
        config = configparser.ConfigParser()
        config.read('config.ini')
        self.logging = int(config['LOGGING']['Logging'])
        self.logging_level = str(config['LOGGING']['LoggingLevel'])
        current_datetime = datetime.now().strftime("%m_%d_%Y_%H_%M_%S")
        logging.basicConfig(filename=f'logs/{current_datetime}_survey_generator.log', filemode='w',
                            format='%(asctime)s - [%(levelname)s]: %(message)s',
                            datefmt='%d-%b-%y %H:%M:%S')
        self.logger = logging.getLogger()
        if self.logging:
            if self.logging_level == 'Debug':
                self.logger.setLevel(logging.DEBUG)
            else:
                self.logger.setLevel(logging.INFO)
        else:
            self.logger.disabled = True
        openai.api_key = config['GPT3 MODEL']['Key']
        self._gpt3_model = config['GPT3 MODEL']['GPT3Model']
        self._chatgpt_model = config['GPT3 MODEL']['ChatGPTModel']
        self._gpt3_temperature = int(
            config['GPT3 MODEL']['GPT3TemperatureDefault'])
        self._gpt3_top_p = int(config['GPT3 MODEL']['GPT3TopP'])
        self._gpt3_freq_penalty = int(
            config['GPT3 MODEL']['GPT3FrequencyPenalty'])
        self._gpt3_presence_penalty = int(
            config['GPT3 MODEL']['GPT3PresencePenalty'])
        
        self._overwrite = int(config['LOGGING']['Overwrite'])
        
        self._max_tokens_business_overview = int(config['GPT3 MODEL']['BusinessOverviewMaxToken'])
        self._temperate_business_overview = float(config['GPT3 MODEL']['BusinessOverviewTemperature'])
        self._max_tokens_research_obj = int(config['GPT3 MODEL']['ResearchObjectivesMaxToken'])
        self._temperate_research_obj = float(config['GPT3 MODEL']['ResearchObjectivesTemperature'])
        self._max_tokens_questionnaire_v2 = int(config['GPT3 MODEL']['QuestionnaireV2MaxToken'])
        self._max_tokens_matrix_oe = int(config['GPT3 MODEL']['MatrixOEMaxToken'])
        self._max_tokens_video_questions = int(config['GPT3 MODEL']['VideoQuestionMaxToken'])
        self._max_tokens_choices_matrix = int(config['GPT3 MODEL']['ChoicesMatrixMaxToken'])
        self._max_tokens_choices_mcq = int(config['GPT3 MODEL']['ChoicesMCQMaxToken'])
        
        self._metric_filename = config['INFERENCE']['MetricsFilename']
        self._min_questions_matrix = int(config['INFERENCE']['MinMatrixQuestions'])
        self._min_questions_matrix_oe = int(config['INFERENCE']['MinMatrixOEQuestions'])
        
        self._prompt_business_overview = str(
            requests.get(config['GPT3 PROMPTS']['BusinessOverview'], timeout=10).text)
        self._prompt_research_obj = str(
            requests.get(config['GPT3 PROMPTS']['ResearchObjectives'], timeout=10).text)
        self._prompt_generate_questionnaire_v2 = str(
            requests.get(config['GPT3 PROMPTS']['GenerateQuestionnaireV2'], timeout=10).text)
        self._prompt_generate_questionnaire_v2_matrix_oe = str(
            requests.get(config['GPT3 PROMPTS']['GenerateQuestionnaireV2MatrixOE'], timeout=10).text)
        self._prompt_generate_video_questions = str(
            requests.get(config['GPT3 PROMPTS']['GenerateVideoQuestions'], timeout=10).text)
        self._prompt_generate_choices_mcq = str(
            requests.get(config['GPT3 PROMPTS']['GenerateChoicesMCQ'], timeout=10).text)
        self._prompt_generate_choices_matrix = str(
            requests.get(config['GPT3 PROMPTS']['GenerateChoicesMatrix'], timeout=10).text)
        
        self._use_chatgpt_business_overview = int(config['LOGGING']['UseChatGPT_business_overview'])
        self._use_chatgpt_research_objective = int(config['LOGGING']['UseChatGPT_research_objective'])
        self._use_chatgpt_survey_generator = int(config['LOGGING']['UseChatGPT_survey_generator'])
        self._use_chatgpt_matrix_oe = int(config['LOGGING']['UseChatGPT_matrix_oe'])
        self._use_chatgpt_video_question = int(config['LOGGING']['UseChatGPT_video_question'])
        self._use_chatgpt_choices_matrix = int(config['LOGGING']['UseChatGPT_choices_matrix'])
        self._use_chatgpt_choices_mcq = int(config['LOGGING']['UseChatGPT_choices_mcq'])

        self.prompt_business_overview_gpt3 = self.read_prompt_json(
            config['OPENAI PROMPTS']['business_overview_gpt3'], False)
        self.prompt_research_objective_gpt3 = self.read_prompt_json(
            config['OPENAI PROMPTS']['research_objective_gpt3'], False)
        self.prompt_survey_generator_gpt3 = self.read_prompt_json(
            config['OPENAI PROMPTS']['survey_generator_gpt3'], False)
        self.prompt_matrix_oe_gpt3 = self.read_prompt_json(
            config['OPENAI PROMPTS']['matrix_oe_gpt3'], False)
        self.prompt_video_question_gpt3 = self.read_prompt_json(
            config['OPENAI PROMPTS']['video_question_gpt3'], False)
        self.prompt_choices_matrix_gpt3 = self.read_prompt_json(
            config['OPENAI PROMPTS']['choices_matrix_gpt3'], False)
        self.prompt_choices_mcq_gpt3 = self.read_prompt_json(
            config['OPENAI PROMPTS']['choices_mcq_gpt3'], False)

        self.prompt_business_overview_chatgpt = self.read_prompt_json(
            config['OPENAI PROMPTS']['business_overview_chatgpt'], True)
        self.prompt_research_objective_chatgpt = self.read_prompt_json(
            config['OPENAI PROMPTS']['research_objective_chatgpt'], True)
        self.prompt_survey_generator_chatgpt = self.read_prompt_json(
            config['OPENAI PROMPTS']['survey_generator_chatgpt'], True)
        self.prompt_matrix_oe_chatgpt = self.read_prompt_json(
            config['OPENAI PROMPTS']['matrix_oe_chatgpt'], True)
        self.prompt_video_question_chatgpt = self.read_prompt_json(
            config['OPENAI PROMPTS']['video_question_chatgpt'], True)
        self.prompt_choices_matrix_chatgpt = self.read_prompt_json(
            config['OPENAI PROMPTS']['choices_matrix_chatgpt'], True)
        self.prompt_choices_mcq_chatgpt = self.read_prompt_json(
            config['OPENAI PROMPTS']['choices_mcq_chatgpt'], True)
        

    @staticmethod
    def read_prompt_json(filename: str, chatgpt=True) -> str:
        """
        Static function to read prompts data from JSON file
        :param filename: Prompt JSON filename
        :return: Prompt string
        """
        if chatgpt:
            with open(f'{filename}', 'r', encoding='utf-8') as file:
                prompt_data = json.load(file)
            return prompt_data['messages']
        with open(f'{filename}', 'r', encoding='utf-8') as file:
            prompt_data = json.load(file)
        return prompt_data['prompt']    

    def get_response_chatgpt(self, prompt: str, max_tokens: int,
                             temperature: float = None) -> (Union[Any, list, dict]):
        """
        Get ChatGPT completion responses for the given prompt
        :param prompt: Prompt to perform completion task on
        :param max_tokens: Max token length for completion
        """
        if not temperature:
            temperature = self._gpt3_temperature
        response = openai.ChatCompletion.create(
            model=self._chatgpt_model,
            messages=prompt,
            temperature=temperature,
            max_tokens=max_tokens,
            top_p=self._gpt3_top_p,
            frequency_penalty=self._gpt3_freq_penalty,
            presence_penalty=self._gpt3_presence_penalty
        )
        time.sleep(20)
        return response        

    def get_response(self, prompt: str, max_tokens: int,
                     temperature: float = None) -> (Union[Any, list, dict]):
        """
        Get GPT3 completion responses for the given prompt
        :param prompt: Prompt to perform completion task on
        :param max_tokens: Max token length for completion
        """
        if not temperature:
            temperature = self._gpt3_temperature
        response = openai.Completion.create(
            model=self._gpt3_model,
            prompt=prompt,
            temperature=temperature,
            max_tokens=max_tokens,
            top_p=self._gpt3_top_p,
            frequency_penalty=self._gpt3_freq_penalty,
            presence_penalty=self._gpt3_presence_penalty
        )
        time.sleep(20)
        return response

    def get_business_overview(self, company_name):
        """
        Function to get the business overview for the given company.
        :param company_name: Name of the company
        """
        self.logger.info('Getting business overview')

        if not self._use_chatgpt_business_overview:
            prompt = self.prompt_business_overview_gpt3.replace("<<COMPANY NAME>>", company_name)
            response = self.get_response(prompt, self._max_tokens_business_overview,
                                        self._temperate_business_overview)
            # Parse the response to get the business overview
            business_overview = response['choices'][0]['text']
            business_overview = company_name + " is" + business_overview
            return business_overview
        prompt = copy.deepcopy(self.prompt_business_overview_chatgpt)
        prompt[1]["content"] = prompt[1]["content"].replace('<<COMPANY NAME>>',company_name)
        response = self.get_response_chatgpt(prompt, self._max_tokens_business_overview,
                                    self._temperate_business_overview)
        # Parse the response to get the business overview
        business_overview = response['choices'][0]['message']['content']
        business_overview = company_name + " is " + business_overview
        time.sleep(20)
        return business_overview

    def get_research_objectives(self, company_name,
                                business_overview,
                                industry, use_case):
        """
        Function to get the research objective for the given company.
        :param company_name: Name of the company
        :param business_overview: Business overview of the company
        :param industry: Industry in which the company operates
        :param use_case: Use Case of the survey
        """
        self.logger.info("Getting research objectives")
        if not self._use_chatgpt_research_objective:
            prompt = self.prompt_research_objective_gpt3.replace("<<COMPANY NAME>>", company_name)
            prompt = prompt.replace("<<INDUSTRY>>", industry)
            prompt = prompt.replace("<<USE CASE>>", use_case)
            prompt = prompt.replace("<<BUSINESS OVERVIEW>>", business_overview)
            response = self.get_response(prompt, self._max_tokens_research_obj,
                                        self._temperate_research_obj)
            research_objectives = response['choices'][0]['text'].strip()
        else:
            prompt = copy.deepcopy(self.prompt_research_objective_chatgpt)
            prompt[1]["content"] = prompt[1]["content"].replace("<<COMPANY NAME>>", company_name).replace("<<INDUSTRY>>", industry).replace("<<USE CASE>>", use_case).replace("<<BUSINESS OVERVIEW>>", business_overview)
            response = self.get_response_chatgpt(prompt, self._max_tokens_research_obj,
                                        self._temperate_business_overview)

            research_objectives = response['choices'][0]['message']['content'].strip()
        ros = research_objectives.split("<")
        l = []
        i = 0
        for r in ros:
            if ">" in r:
                s = r.split(">")
                if i == 0:
                    s[0] = '<mark style="background-color: blanchedalmond;">' + s[0] + '</mark>'
                elif i == 1:
                    s[0] = '<mark style="background-color: aqua;">' + s[0] + '</mark>'
                elif i == 2:
                    s[0] = '<mark style="background-color: #90ee90;">' + s[0] + '</mark>'
                t = "".join(s)
                l.append(t)
                i+=1
            else:
                l.append(r)
        research_objectives = "".join(l)
        time.sleep(20)
        return research_objectives

    def get_questionnaire(self, company_name, business_overview,
                          research_objectives) -> Tuple[str, List[str]]:
        """
        Function to get the Questionnaire.
        :param company_name: Name of the company
        :param business_overview: Business overview of the company
        :param research_objectives: Survey research objective
        """
        if not self._use_chatgpt_survey_generator:
            prompt = self.prompt_survey_generator_gpt3.replace("<<BUSINESS OVERVIEW>>", business_overview)
            prompt = prompt.replace("<<RESEARCH OBJECTIVES>>", research_objectives)
            prompt = prompt.replace("<<COMPANY NAME>>", company_name)
            response = self.get_response(prompt, self._max_tokens_questionnaire_v2)
            questionnaire_string = response['choices'][0]['text']
            questions = questionnaire_string.split("\n")
            self.logger.info("---------------- FIRST GENERATION ----------------")
            self.logger.info(questionnaire_string)
            return questionnaire_string, questions
        prompt = copy.deepcopy(self.prompt_survey_generator_chatgpt)
        prompt[0]["content"] = prompt[0]["content"].replace("<<COMPANY NAME>>", company_name)
        prompt[1]["content"] = prompt[1]["content"].replace("<<BUSINESS OVERVIEW>>", business_overview).replace("<<RESEARCH OBJECTIVES>>", research_objectives).replace("<<COMPANY NAME>>", company_name)
        response = self.get_response_chatgpt(prompt, self._max_tokens_questionnaire_v2)
        questionnaire_string = response['choices'][0]['message']['content']
        questions = questionnaire_string.split("\n")
        self.logger.info("---------------- FIRST GENERATION ----------------")
        self.logger.info(questionnaire_string)
        time.sleep(20)
        return questionnaire_string, questions

    def get_matrix_questions(self, questionnaire_string, company_name, business_overview,
                             research_objectives) -> str:
        """
        Function to get the Matrix Questions.
        :param questionnaire_string: Previously generated Questionnaire string
        :param company_name: Name of the company
        :param business_overview: Business overview of the company
        :param research_objectives: Survey research objective
        :return: Updated Questionnaire string
        """
        last_list_number = int(questionnaire_string.split("\n")[-1].split(".")[0])
        if not self._use_chatgpt_matrix_oe: 
            prompt = self.prompt_matrix_oe_gpt3.replace("<<BUSINESS OVERVIEW>>", business_overview)
            prompt = prompt.replace("<<RESEARCH OBJECTIVES>>", research_objectives)
            prompt = prompt.replace("<<COMPANY NAME>>", company_name)
            prompt = prompt + questionnaire_string + "\n" + str(last_list_number + 1) + ". [Matrix]"
            response = self.get_response(prompt, self._max_tokens_matrix_oe)
            questionnaire_string = questionnaire_string + "\n" + str(last_list_number + 1) + ". [Matrix]" + response['choices'][0]['text']
            self.logger.info("---------------- MATRIX Q GENERATION ----------------")
            self.logger.info(questionnaire_string)
            return questionnaire_string
        self.logger.info('Getting matrix question chatgpt')
        prompt = copy.deepcopy(self.prompt_matrix_oe_chatgpt)
        prompt[0]["content"] = prompt[0]["content"].replace("<<COMPANY NAME>>", company_name)
        prompt[1]["content"] = prompt[1]["content"].replace("<<BUSINESS OVERVIEW>>", business_overview).replace("<<RESEARCH OBJECTIVES>>", research_objectives).replace("<<COMPANY NAME>>", company_name)
        prompt[1]["content"] = prompt[1]["content"] + questionnaire_string + "\n" + str(last_list_number + 1) + ". [Matrix]"
        response = self.get_response_chatgpt(prompt, self._max_tokens_matrix_oe)
        questionnaire_string = questionnaire_string + "\n" + str(last_list_number + 1) + ". [Matrix]" + response['choices'][0]['message']['content']
        self.logger.info("---------------- MATRIX Q GENERATION ----------------")
        self.logger.info(questionnaire_string)
        time.sleep(20)
        return questionnaire_string

    def get_open_ended_questions(self, questionnaire_string, company_name,
                                business_overview, research_objectives) -> str:
        """
        Function to get the Matrix Open-ended Questions.
        :param questionnaire_string: Previously generated Questionnaire string
        :param company_name: Name of the company
        :param business_overview: Business overview of the company
        :param research_objectives: Survey research objective
        :return: Updated Questionnaire string
        """
        # Find the last list number in the questionnaire string
        last_list_number = int(questionnaire_string.split("\n")[-1].split(".")[0])
        if not self._use_chatgpt_matrix_oe: 
            prompt = self.prompt_matrix_oe_gpt3.replace("<<BUSINESS OVERVIEW>>", business_overview)
            prompt = prompt.replace("<<RESEARCH OBJECTIVES>>", research_objectives)
            prompt = prompt.replace("<<COMPANY NAME>>", company_name)
            prompt = prompt + questionnaire_string + "\n" + str(last_list_number + 1) + ". [Open-ended]"
            response = self.get_response(prompt, self._max_tokens_matrix_oe)
            questionnaire_string = questionnaire_string + "\n" + str(last_list_number + 1) + ". [Open-ended]" + response['choices'][0]['text']
            self.logger.info("---------------- OE Q GENERATION ----------------")
            self.logger.info(questionnaire_string)
            return questionnaire_string
        prompt = copy.deepcopy(self.prompt_matrix_oe_chatgpt)
        prompt[0]["content"] = prompt[0]["content"].replace("<<COMPANY NAME>>", company_name)
        prompt[1]["content"] = prompt[1]["content"].replace("<<BUSINESS OVERVIEW>>", business_overview).replace("<<RESEARCH OBJECTIVES>>", research_objectives).replace("<<COMPANY NAME>>", company_name)
        prompt[1]["content"] = prompt[1]["content"] + questionnaire_string + "\n" + str(last_list_number + 1) + ". [Open-ended]"
        response = self.get_response_chatgpt(prompt, self._max_tokens_matrix_oe)
        questionnaire_string = questionnaire_string + "\n" + str(last_list_number + 1) + ". [Open-ended]" + response['choices'][0]['message']['content']
        self.logger.info("---------------- OE Q GENERATION ----------------")
        self.logger.info(questionnaire_string)
        time.sleep(20)
        return questionnaire_string

    def get_video_questions(self, company_name,
                            business_overview, research_objectives) -> list:
        """
        Function to get the Video Questions.
        :param questionnaire_string: Previously generated Questionnaire string
        :param company_name: Name of the company
        :param business_overview: Business overview of the company
        :param research_objectives: Survey research objective
        """
        if not self._use_chatgpt_video_question:
            prompt = self.prompt_video_question_gpt3.replace("<<BUSINESS OVERVIEW>>", business_overview)
            prompt = prompt.replace("<<RESEARCH OBJECTIVES>>", research_objectives)
            prompt = prompt.replace("<<COMPANY NAME>>", company_name)
            response = self.get_response(prompt, self._max_tokens_video_questions)
            video_questions = response['choices'][0]['text'].split('\n')
            self.logger.info("---------------- VIDEO Q GENERATION ----------------")
            self.logger.info(video_questions)
            return video_questions
        prompt = copy.deepcopy(self.prompt_video_question_chatgpt)
        prompt[1]["content"] = prompt[1]["content"].replace("<<BUSINESS OVERVIEW>>", business_overview).replace("<<RESEARCH OBJECTIVES>>", research_objectives).replace("<<COMPANY NAME>>", company_name)
        response = self.get_response_chatgpt(prompt, self._max_tokens_video_questions)
        video_questions = response['choices'][0]['message']['content'].split('\n')
        self.logger.info("---------------- VIDEO Q GENERATION ----------------")
        self.logger.info(video_questions)
        time.sleep(20)
        return video_questions

    def get_choices_matrix(self, question, company_name,
                           business_overview, research_objectives):
        """
        Function to get Choices for Matrix Question.
        :param question: Matrix Question
        :param company_name: Name of the company
        :param business_overview: Business overview of the company
        :param research_objectives: Survey research objective
        """
        if not self._use_chatgpt_choices_matrix:
            prompt = self.prompt_choices_matrix_gpt3.replace("<<QUESTION>>", question)
            prompt = prompt.replace("<<BUSINESS OVERVIEW>>", business_overview)
            prompt = prompt.replace("<<RESEARCH OBJECTIVES>>", research_objectives)
            prompt = prompt.replace("<<COMPANY NAME>>", company_name)
            response = self.get_response(prompt, self._max_tokens_choices_matrix)
            matrix_choices = response['choices'][0]['text'].split("Columns:")
            rows = matrix_choices[0].split("\n")
            columns = matrix_choices[1].split("\n")
            return rows, columns
        prompt = copy.deepcopy(self.prompt_choices_matrix_chatgpt)
        prompt[0]["content"] = prompt[0]["content"].replace("<<COMPANY NAME>>", company_name)
        prompt[1]["content"] = prompt[1]["content"].replace("<<BUSINESS OVERVIEW>>", business_overview).replace("<<RESEARCH OBJECTIVES>>", research_objectives).replace("<<COMPANY NAME>>", company_name).replace("<<QUESTION>>", question)
        response = self.get_response_chatgpt(prompt, self._max_tokens_choices_matrix)
        matrix_choices = response['choices'][0]['message']['content'].split("Columns:")
        rows = matrix_choices[0].split("\n")
        columns = matrix_choices[1].split("\n")
        time.sleep(20)
        return rows, columns

    def get_choices_mcq(self, question, company_name,
                        business_overview, research_objectives):
        """
        Function to get Choices for MCQ.
        :param question: Matrix Question
        :param company_name: Name of the company
        :param business_overview: Business overview of the company
        :param research_objectives: Survey research objective
        """
        if not self._use_chatgpt_choices_mcq:
            prompt = self.prompt_choices_mcq_gpt3.replace("<<QUESTION>>", question)
            prompt = prompt.replace("<<BUSINESS OVERVIEW>>", business_overview)
            prompt = prompt.replace("<<RESEARCH OBJECTIVES>>", research_objectives)
            prompt = prompt.replace("<<COMPANY NAME>>", company_name)
            response = self.get_response(prompt, self._max_tokens_choices_matrix)
            choices = response['choices'][0]['text'].split("\n")
            return choices
        prompt = copy.deepcopy(self.prompt_choices_mcq_chatgpt)
        prompt[0]["content"] = prompt[0]["content"].replace("<<COMPANY NAME>>", company_name)
        prompt[1]["content"] = prompt[1]["content"].replace("<<BUSINESS OVERVIEW>>", business_overview).replace("<<RESEARCH OBJECTIVES>>", research_objectives).replace("<<COMPANY NAME>>", company_name).replace("<<QUESTION>>", question)
        response = self.get_response_chatgpt(prompt, self._max_tokens_choices_matrix)
        choices = response['choices'][0]['message']['content'].split("\n")
        time.sleep(20)
        return choices

    def surveyjs_questionnaire(self, questionnaire: list) -> list:
        """
        Function to convert questionnaire to surveyjs_questionnaire.
        :param questionnaire: Questionnaire list
        :return surveyjs_questionnaire list
        """
        pages = []
        for page_idx, question in enumerate(questionnaire, start=1):
            if question['type'] == "Multiple Choice" or question['type'] == "Multiple choice":
                text = question['question'].lower()
                if 'select all' in text:
                    q_type = 'checkbox'
                else:
                    q_type = 'radiogroup'
                choices = []
                for choice in question['choices']:
                    choice_dict = {
                        "value": choice,
                        "text": f'<p>{choice}</p>'
                    }
                    choices.append(choice_dict)
                js_dict = {
                "name": f"page{page_idx}",
                "elements": [
                    {
                    "type": q_type,
                    "name": f"question{page_idx}",
                    "title": f"<p>{question['question']}</p>",
                    "surveyQID": str(uuid.uuid1()),
                    "choices": choices
                    }
                ]
                }
            elif question['type'] == 'Open-ended':
                js_dict = {
                "name": f"page{page_idx}",
                "elements": [
                    {
                    "type": "comment",
                    "name": f"question{page_idx}",
                    "title": f"<p>{question['question']}</p>",
                    "surveyQID": str(uuid.uuid1())
                    }
                ]
                }
            elif question['type'] == 'Matrix':
                columns = []
                rows = []
                for choice in question['choices'][0]:
                    row_dict = {
                        "value": choice,
                        "text": f'<p>{choice}</p>'
                    }
                    rows.append(row_dict)
                for choice in question['choices'][1]:
                    col_dict = {
                        "value": choice,
                        "text": f'<p>{choice}</p>'
                    }
                    columns.append(col_dict)
                js_dict = {
                "name": f"page{page_idx}",
                "elements": [
                    {
                    "type": "matrix",
                    "name": f"question{page_idx}",
                    "title": f"<p>{question['question']}</p>",
                    "surveyQID": str(uuid.uuid1()),
                    "columns": columns,
                    "rows": rows 
                    }
                ]
                }
            else:
                js_dict = {
                "name": f"page{page_idx}",
                "elements": [
                    {
                    "type": "videofeedback",
                    "name": f"question{page_idx}",
                    "title": f"<p>{question['question']}</p>",
                    "surveyQID": str(uuid.uuid1())
                    }
                ]
                }
            pages.append(js_dict)
        return pages

    def create_survey(self, company_name, business_overview,
                      research_objectives, project_name, request_id=0):
        """
        Function to create a questionnaire containing multiple questions and choices for
        those questions for the user based on the research objectives and business overview.
        :param company_name: Name of the company
        :param business_overview: Business overview of the company
        :param research_objectives: Survey research objective
        :param project_name: Survey project name
        """
        # If questionnaire.json exists, return the questionnaire from the file
        if os.path.exists(f'questionnaires/questionnaire_{project_name.replace(" ", "_").replace("/", "-")}.json'):
            self.logger.info('Using saved Quesstionaire JSON!!!')
            with open(f'questionnaires/questionnaire_{project_name.replace(" ", "_").replace("/", "-")}.json') as json_file:
                questionnaire = json.load(json_file)
                questionnaire_js = self.surveyjs_questionnaire(questionnaire)
            return questionnaire, questionnaire_js
        questionnaire = []
        # Use the openai api to generate a questionnaire based on the
        # research objectives and business overview
        questionnaire_string, questions = self.get_questionnaire(company_name,
                                                                business_overview,
                                                                research_objectives)

        # Dump the questionnaire string into a txt file for debugging purposes
        with open('questionnaire.txt', 'w') as file:
            file.write(questionnaire_string)

        # Figure out the number of Matrix questions
        matrix_questions = 0
        for question in questions:
            if "[Matrix]" in question:
                matrix_questions += 1
        # Figure out the number of open-ended questions
        open_ended_questions = 0
        for question in questions:
            if "[Open-ended]" in question:
                open_ended_questions += 1
        # Figure out the number of multiple choice questions
        multiple_choice_questions = 0
        for question in questions:
            if "[Multiple Choice]" in question:
                multiple_choice_questions += 1

        # We want to generate at least 3 matrix questions
        while matrix_questions < self._min_questions_matrix:
            questionnaire_string = self.get_matrix_questions(questionnaire_string, company_name, 
                                                             business_overview, research_objectives)
            matrix_questions += 1
        # We want to generate at least 3 open-ended questions
        while open_ended_questions < self._min_questions_matrix_oe:
            questionnaire_string = self.get_open_ended_questions(questionnaire_string,
                                                                 company_name,
                                                                 business_overview,
                                                                 research_objectives)
            open_ended_questions += 1

        # Split the questionnaire string into a list of question dictionaries
        questions = questionnaire_string.split("\n")
        
        print(f"########### question:\n {questions} \n###############")
        # Iterate through the questions and split each question into the question and the choices
        for question in questions:
            # Convert each question into a dictionary with the question, type, and choices
            question_dict = {}
            question_dict['question'] = question.split("]")[0].strip()
            question_dict['type'] = question.split("[")[0].split("]")[0].strip()
            question_dict['choices'] = []
            questionnaire.append(question_dict)

        # Generate video questions
        video_questions = self.get_video_questions(company_name, business_overview,
                                                   research_objectives)
        for question in video_questions:
            if not question.strip():  # Skip empty strings
                continue
            if question[0] == "2" or question[0] == "3":
                question = question[3:]
            question_dict = {}
            question_dict['question'] = question.strip()
            question_dict['type'] = "Video"
            question_dict['choices'] = ['']
            questionnaire.append(question_dict)

        # Generate choices for each question
        for question in questionnaire:
            self.logger.info(f"Generating choices for question: {question['question']}")
            # If the question is a matrix question, generate choices for the matrix
            if question['type'] == "Matrix":
                question['choices'] = [[], []]
                question['question'] = question['question']
                rows, columns = self.get_choices_matrix(question['question'], company_name,
                                                        business_overview, research_objectives)
                for row in rows:
                    if row != "":
                        row = row.replace("-", "", 1).strip()
                        question['choices'][0].append(row)
                for column in columns:
                    if column != "":
                        column = column.replace("-", "", 1).strip()
                        question['choices'][1].append(column)

            # If the question is a multiple choice question, generate choices for
            # the multiple choice question
            if question['type'] == "Multiple Choice" or question['type'] == "Multiple choice":
                choices = self.get_choices_mcq(question['question'],
                                               company_name,
                                               business_overview,
                                               research_objectives)
                for choice in choices:
                    choice = choice.replace('-', '', 1).strip()
                    question['choices'].append(choice)

            # if the question is an open-ended question, just write
            # "Open-ended text response" as the only choice
            if question['type'] == "Open-ended":
                question['choices'] = ["Open-ended text response"]
        self.logger.info("Completed choice generation for questions!!!")

        # remove questions from the questionnaire that have the words "gender" or
        # "ethnicity" or "your age" or "how old"
        new_questionnaire = []
        for question in questionnaire:
            if 'gender' not in question['question'].lower()\
            or 'ethnicity' not in question['question'].lower()\
            or 'your age' not in question['question'].lower()\
            or 'how old' not in question['question'].lower():
                new_questionnaire.append(question)
        questionnaire = new_questionnaire
        # questionnaire = [{'question': 'In which state and city do you live?', 'type': 'Multiple Choice', 'choices' : ['Dropdown of States', 'Dropdown of Cities']}] + questionnaire
        # questionnaire = [{'question': 'What is your race/ethnicity?', 'type': 'Multiple Choice', 'choices' : ["White", 'Black or African American', 'Hispanic or Latino', 'Asian', 'American Indian or Alaska Native', 'Native Hawaiian or Other Pacific Islander', 'Other', 'Prefer not to say']}] + questionnaire
        # questionnaire = [{'question': 'What is your Gender?', 'type': 'Multiple Choice', 'choices' : ["Male", 'Female', 'Other', 'Prefer not to say']}] + questionnaire
        # questionnaire = [{'question': 'What year were you born?', 'type': 'Multiple Choice', 'choices': ['Dropdown of Years']}] + questionnaire

        # Remove duplicate questions
        questions_list = []
        new_questionnaire = []
        for question in questionnaire:
            if question['question'].lower() in questions_list:
                continue
            questions_list.append(question['question'])
            new_questionnaire.append(question)
        questionnaire = new_questionnaire

        # Convert questionnaire to Survey JS format dictionary
        questionnaire_js = self.surveyjs_questionnaire(questionnaire)

        
        # Save the questionnaire to a json file
        with open(f'questionnaires/questionnaire_{project_name.replace(" ", "_").replace("/", "-")}_{request_id}.json', 'w') as outfile:
            json.dump(questionnaire, outfile, indent=2)
        return questionnaire, questionnaire_js

    def export_docx(self, project_name, company_name,
                    research_objectives, questionnaire, request_id=0):
        """
        Function to export the Questionnaire to Docx file and return link
        :param company_name: Name of the company
        :param research_objectives: Survey research objective
        :param project_name: Survey project name
        """
        self.logger.info('Exporting to Docx file!!!')
        # Open the template docx file
        template = Document("template_new.docx")

        for paragraph in template.paragraphs:
            paragraph.text = paragraph.text.replace('<<PROJECT NAME>>', project_name)
            paragraph.text = paragraph.text.replace('<<COMPANY>>', company_name)
            paragraph.text = paragraph.text.replace('<<RESEARCH OBJECTIVES>>', research_objectives)
            #paragraph.text = paragraph.text.replace('<<Request ID>>', str(request_id))
            self.logger.info(f'template :{paragraph.text}')
        
        # Insert the questionnaire into the template as a numbered list
        self.logger.info('started question')
        for question in questionnaire:
            # Add the question to the template
            template.add_paragraph(question['question'])
            # Add the choices to the template
            for choice_idx, choice in enumerate(question['choices']):
                if isinstance(choice, list):
                    self.logger.info(question)
                    self.logger.info(choice_idx)
                    self.logger.info(choice)
                    if choice_idx == 0:
                        template.add_paragraph("Rows: ")
                        for c in choice:
                            template.add_paragraph(str(c), style='List Bullet')
                    if choice_idx == 1:
                        template.add_paragraph("Columns: ")
                        for c in choice:
                            template.add_paragraph(str(c), style='List Bullet')
                        template.add_paragraph()
                else:
                    template.add_paragraph(choice, style='List Bullet')
            # Add a blank line
            template.add_paragraph()
        self.logger.info('ended question')
        # Save the template as a new docx file
        template.save(f"questionnaires/{project_name.replace(' ', '_')}_questionnaire_{request_id}.docx")
        self.logger.info('ended saving')
        files = {
                'files[]': open(f"questionnaires/{project_name.replace(' ', '_')}_questionnaire_{request_id}.docx", 'rb'),
            }

        response = requests.post('https://uguu.se/upload', files=files, timeout=60)

        # Read the response as json
        self.logger.info(response.json())

        if response.json()['success']:
            # Get the download link
            self.logger.info(response.json()['files'][0]['url'])
            download_link = response.json()['files'][0]['url']
            download_link = "https://docs.google.com/viewerng/viewer?url=" + download_link
        else:
            download_link = "Error"
        return download_link

    def update_metrics(self, project_name, company_name, questionnaire: List[dict], execution_time) -> None:
        """
        Function to update the metrics for a project.
        :param project_name: Name of the project
        :param company_name: Name of the company
        :param questionnaire: List of questions dictionary containing question, type and choices
        """
        metrics_data = {
            "Project Name": project_name,
            "Company Name": company_name,
            "Execution time": execution_time,
            "No. of Multiple Choice questions": 0,
            "No. of Open-ended questions": 0,
            "No. of Matrix questions": 0,
            "No. of Video questions": 0,
        }
        for question_dict in questionnaire:
            metrics_data[f'No. of {question_dict["type"]} questions'] += 1
        metrics_df = pd.Series(metrics_data).to_frame().T
        if not os.path.exists(self._metric_filename):
            metrics_df.to_csv(self._metric_filename, index=False)
        else:
            metrics_df.to_csv(self._metric_filename, header=False,
                              index=False, mode='a')
    